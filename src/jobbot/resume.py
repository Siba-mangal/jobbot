"""Resume extraction.

The resume is the stable half of every scoring prompt, so it gets extracted
once and cached to disk. Byte-stability matters: the cached text becomes the
cached prompt prefix, and any change invalidates it.
"""

from __future__ import annotations

import hashlib
import re
from pathlib import Path

from .config import DATA_DIR

# Dotted, distinct names: a user's own resume may legitimately be
# `data/resume.txt`, and the cache must never be able to overwrite the source
# document it was extracted from.
CACHE_PATH = DATA_DIR / ".resume-cache.txt"
HASH_PATH = DATA_DIR / ".resume-cache.sha256"

_WS_RUN = re.compile(r"[ \t]+")
_BLANK_RUN = re.compile(r"\n{3,}")


class ResumeError(RuntimeError):
    pass


def _extract_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = [(page.extract_text() or "") for page in reader.pages]
    return "\n\n".join(pages)


def _extract_docx(path: Path) -> str:
    import docx

    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def _clean(text: str) -> str:
    """Normalize whitespace so the same PDF always yields the same bytes."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = "".join(ch for ch in text if ch == "\n" or ch == "\t" or ch >= " ")
    lines = [_WS_RUN.sub(" ", line).rstrip() for line in text.split("\n")]
    return _BLANK_RUN.sub("\n\n", "\n".join(lines)).strip()


def extract_text(path: Path) -> str:
    if not path.exists():
        raise ResumeError(f"Resume not found at {path}")
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        raw = _extract_pdf(path)
    elif suffix in {".docx", ".doc"}:
        raw = _extract_docx(path)
    elif suffix in {".txt", ".md"}:
        raw = path.read_text(errors="replace")
    else:
        raise ResumeError(
            f"Unsupported resume format {suffix!r}. Use PDF, DOCX, TXT, or MD."
        )

    text = _clean(raw)
    if len(text) < 200:
        raise ResumeError(
            f"Extracted only {len(text)} characters from {path.name}. "
            "The file may be a scanned image rather than text — export a text-based "
            "PDF, or save a .txt copy alongside it and point profile.yaml at that."
        )
    return text


def _source_hash(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def get_resume_text(path: Path, *, refresh: bool = False) -> str:
    """Extracted resume text, cached until the source file changes.

    Cache is keyed on the source file's hash, so replacing your resume
    automatically re-extracts (and correctly invalidates the prompt cache
    downstream, since the prefix bytes change).
    """
    current = _source_hash(path)
    if not refresh and CACHE_PATH.exists() and HASH_PATH.exists():
        if HASH_PATH.read_text().strip() == current:
            return CACHE_PATH.read_text()

    text = extract_text(path)
    CACHE_PATH.parent.mkdir(parents=True, exist_ok=True)
    CACHE_PATH.write_text(text)
    HASH_PATH.write_text(current)
    return text
